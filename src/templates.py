"""
Contract Templates with AI-Powered Population

Provides structured legal templates with placeholders that can be dynamically
populated from extracted contract data. Supports both .txt and .docx formats.
"""

import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Template structures with consistent placeholder naming
SERVICE_AGREEMENT_TEMPLATE = """SERVICE AGREEMENT

This Service Agreement ("Agreement") is made and entered into on [START_DATE], by and between:

PROVIDER: [PROVIDER] ("Provider")
CLIENT: [CLIENT] ("Client")

1. SERVICES
The Provider agrees to provide the following services:
[SERVICES]

2. COMPENSATION
Total Amount: [AMOUNT]
Payment Terms: [PAYMENT_TERMS]

3. TERM AND TERMINATION
This Agreement shall commence on [START_DATE] and continue until [END_DATE].
Termination Notice: [TERMINATION_NOTICE]

4. CONFIDENTIALITY
[CONFIDENTIALITY]

5. GOVERNING LAW
This Agreement shall be governed by the laws of [JURISDICTION].

IN WITNESS WHEREOF, the parties have executed this Agreement on the date first above written.

______________________                  ______________________
Provider Signature                      Client Signature
"""

NDA_TEMPLATE = """NON-DISCLOSURE AGREEMENT (NDA)

This Non-Disclosure Agreement ("Agreement") is entered into on [START_DATE] between:

DISCLOSING PARTY: [PROVIDER]
RECEIVING PARTY: [CLIENT]

1. PURPOSE
The parties wish to explore a business opportunity and protect confidential information.

2. CONFIDENTIAL INFORMATION
"Confidential Information" means all non-public information disclosed by one party to the other, including but not limited to:
[CONFIDENTIALITY]

3. OBLIGATIONS
The Receiving Party agrees to:
a) Use Confidential Information only for the stated Purpose
b) Not disclose such information to any third party without prior written consent
c) Protect the information with the same degree of care used for own confidential information

4. TERM
The obligations of this Agreement shall survive for a period of [TERMINATION_NOTICE].

5. JURISDICTION
This Agreement shall be governed by the laws of [JURISDICTION].

______________________                  ______________________
Disclosing Party                        Receiving Party
"""

EMPLOYMENT_TEMPLATE = """EMPLOYMENT AGREEMENT

Date: [START_DATE]

EMPLOYER: [PROVIDER]
EMPLOYEE: [CLIENT]

1. POSITION
The Employee is appointed to the position described as follows:
[SERVICES]

2. COMPENSATION
Annual Compensation: [AMOUNT]
Payment Terms: [PAYMENT_TERMS]

3. EMPLOYMENT PERIOD
Start Date: [START_DATE]
End Date: [END_DATE]

4. TERMINATION
Notice Period: [TERMINATION_NOTICE]

5. CONFIDENTIALITY
[CONFIDENTIALITY]

6. GOVERNING LAW
This Agreement shall be governed by the laws of [JURISDICTION].

For Employer:                           Employee Acknowledgment:

______________________                  ______________________
Authorized Signatory                    Employee Signature
"""


def get_template_structure(template_type):
    """
    Get the raw template structure for a given template type.
    
    Args:
        template_type: Name of the template ("Service Agreement", "NDA", "Employment Agreement")
        
    Returns:
        Template string with placeholders
    """
    templates = {
        "Service Agreement": SERVICE_AGREEMENT_TEMPLATE,
        "NDA": NDA_TEMPLATE,
        "Employment Agreement": EMPLOYMENT_TEMPLATE
    }
    return templates.get(template_type, "")


def populate_template(template_type, extracted_data):
    """
    Populate a template with extracted contract data.
    
    Args:
        template_type: Name of the template to populate
        extracted_data: Dictionary of extracted fields from contract
        
    Returns:
        Populated template string with all placeholders replaced
    """
    template = get_template_structure(template_type)
    
    if not template:
        return f"Error: Template '{template_type}' not found."
    
    # Create a copy of extracted data with uppercase keys for placeholder matching
    data_map = {key.upper(): value for key, value in extracted_data.items()}
    
    # Replace all placeholders with extracted data
    for key, value in data_map.items():
        placeholder = f"[{key}]"
        template = template.replace(placeholder, str(value))
    
    # Replace any remaining placeholders with fallback text
    template = re.sub(r'\[([^\]]+)\]', 'Not specified in the provided contract', template)
    
    return template


def generate_docx(template_type, extracted_data):
    """
    Generate a Word document (.docx) from populated template.
    
    Args:
        template_type: Name of the template
        extracted_data: Dictionary of extracted fields
        
    Returns:
        BytesIO buffer containing the .docx file
    """
    # Get populated text
    populated_text = populate_template(template_type, extracted_data)
    
    # Create Word document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph(template_type.upper())
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    # Add content
    lines = populated_text.split('\n')
    for line in lines[1:]:  # Skip first line (title already added)
        if line.strip():
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            p.paragraph_format.line_spacing = 1.15
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def get_templates():
    """
    Get dictionary of all available templates (legacy function for compatibility).
    
    Returns:
        Dictionary mapping template names to template structures
    """
    return {
        "Service Agreement": SERVICE_AGREEMENT_TEMPLATE,
        "Non-Disclosure Agreement (NDA)": NDA_TEMPLATE,
        "Employment Contract": EMPLOYMENT_TEMPLATE
    }
