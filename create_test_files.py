"""
Generate test contract files in specific formats for testing.
Creates exactly 5 files: 3 PDFs, 1 DOCX, 1 TXT
"""

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, Inches
import os

# Create test_contracts directory
os.makedirs("test_contracts", exist_ok=True)

# 1. SERVICE AGREEMENT (PDF)
def create_service_agreement_pdf():
    filename = "test_contracts/service_agreement.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, alignment=1)
    story.append(Paragraph("SERVICE AGREEMENT", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    content = """
This Service Agreement ("Agreement") is made and entered into on January 15, 2024, by and between:<br/><br/>

<b>PROVIDER:</b> TechCorp Solutions Ltd., a company incorporated under the laws of Delaware, USA, with its principal place of business at 123 Innovation Drive, San Francisco, CA 94105 ("Provider")<br/><br/>

<b>AND</b><br/><br/>

<b>CLIENT:</b> ABC Enterprises Inc., a company incorporated under the laws of California, USA, with its principal place of business at 456 Business Park, Los Angeles, CA 90001 ("Client")<br/><br/>

<b>1. SERVICES</b><br/>
The Provider agrees to provide the following services to the Client:<br/>
- Development of a mobile application for iOS and Android platforms<br/>
- UI/UX design and user interface development<br/>
- Backend API integration with Client's existing systems<br/>
- Payment gateway integration (Stripe and PayPal)<br/>
- Deployment to Apple App Store and Google Play Store<br/>
- Post-launch support for 30 days<br/><br/>

<b>2. COMPENSATION</b><br/>
2.1 Total Fee: The Client agrees to pay the Provider a total fee of Fifty Thousand US Dollars ($50,000 USD).<br/><br/>

2.2 Payment Schedule:<br/>
- 30% ($15,000) upon signing this Agreement<br/>
- 40% ($20,000) upon completion of development and testing phase<br/>
- 30% ($15,000) upon final delivery and deployment to app stores<br/><br/>

2.3 Payment Terms: All payments shall be made within 90 days after invoice date via wire transfer.<br/><br/>

<b>3. TERM AND TERMINATION</b><br/>
3.1 Term: This Agreement shall commence on January 15, 2024, and shall continue until December 31, 2024.<br/><br/>

3.2 Termination by Provider: The Provider may terminate this Agreement with 24 hours written notice if the Client fails to make any payment when due.<br/><br/>

3.3 Termination by Client: The Client may terminate this Agreement by providing 6 months written notice to the Provider.<br/><br/>

<b>4. CONFIDENTIALITY</b><br/>
Both parties acknowledge that they may have access to confidential information during the term of this Agreement and agree to maintain the confidentiality of such information.<br/><br/>

<b>5. NON-COMPETE</b><br/>
The Client agrees not to hire, solicit, or engage any of the Provider's employees or contractors for a period of 5 years from the date of this Agreement.<br/><br/>

<b>6. LIMITATION OF LIABILITY</b><br/>
The Provider's total liability under this Agreement shall be limited to the total fees paid by the Client.<br/><br/>

<b>7. GOVERNING LAW</b><br/>
This Agreement shall be governed by and construed in accordance with the laws of the State of California, USA.<br/><br/>

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.<br/><br/>

PROVIDER: TechCorp Solutions Ltd.<br/>
By: John Smith, CEO<br/>
Date: January 15, 2024<br/><br/>

CLIENT: ABC Enterprises Inc.<br/>
By: Sarah Johnson, CTO<br/>
Date: January 15, 2024
"""
    
    story.append(Paragraph(content, styles['Normal']))
    doc.build(story)
    print(f"[OK] Created: {filename}")

# 2. EMPLOYMENT CONTRACT (DOCX)
def create_employment_contract_docx():
    filename = "test_contracts/employment_contract.docx"
    doc = Document()
    
    # Title
    title = doc.add_heading('EMPLOYMENT AGREEMENT', 0)
    title.alignment = 1  # Center
    
    doc.add_paragraph()
    
    content = [
        ("This Employment Agreement (\"Agreement\") is entered into on March 1, 2024, between:", False),
        ("", False),
        ("EMPLOYER: InnovateTech Private Limited", True),
        ("Registered Office: Tower B, Tech Park, Whitefield, Bangalore - 560066, Karnataka, India", False),
        ("CIN: U72900KA2015PTC123456", False),
        ("(\"Company\" or \"Employer\")", False),
        ("", False),
        ("AND", False),
        ("", False),
        ("EMPLOYEE: Mr. Rajesh Kumar", True),
        ("Address: Flat 301, Green Meadows Apartment, HSR Layout, Bangalore - 560102, Karnataka, India", False),
        ("(\"Employee\")", False),
        ("", False),
        ("1. POSITION AND DUTIES", True),
        ("1.1 Position: The Employee is appointed to the position of Software Engineer - Full Stack Development.", False),
        ("", False),
        ("1.2 Duties: The Employee shall develop and maintain web applications using React, Node.js, and MongoDB, participate in code reviews, collaborate with cross-functional teams, and mentor junior developers.", False),
        ("", False),
        ("2. COMPENSATION AND BENEFITS", True),
        ("2.1 Annual Salary: The Employee shall receive an annual Cost to Company (CTC) of Twelve Lakh Rupees (₹12,00,000) per annum.", False),
        ("", False),
        ("2.2 Salary Structure:", False),
        ("- Basic Salary: ₹6,00,000 per annum", False),
        ("- House Rent Allowance: ₹2,40,000 per annum", False),
        ("- Special Allowance: ₹1,80,000 per annum", False),
        ("- Employer's Provident Fund Contribution: ₹72,000 per annum", False),
        ("", False),
        ("2.3 Payment Schedule: Salary shall be paid monthly on the last working day of each calendar month.", False),
        ("", False),
        ("3. EMPLOYMENT PERIOD", True),
        ("3.1 Commencement: Employment shall commence on March 1, 2024.", False),
        ("3.2 Probation: The Employee shall be on probation for 6 months. During probation, either party may terminate with 15 days notice.", False),
        ("", False),
        ("4. TERMINATION", True),
        ("4.1 Notice Period: After confirmation, either party may terminate this Agreement by giving 60 days written notice.", False),
        ("", False),
        ("5. CONFIDENTIALITY AND NON-DISCLOSURE", True),
        ("5.1 The Employee acknowledges access to Confidential Information including source code, algorithms, business strategies, and customer information.", False),
        ("5.2 The Employee agrees to maintain strict confidentiality and not disclose such information to any third party.", False),
        ("5.3 These confidentiality obligations shall survive termination for 3 years.", False),
        ("", False),
        ("6. NON-COMPETE", True),
        ("6.1 During employment and for 12 months after termination, the Employee shall not engage in any competing business or solicit Company customers or employees.", False),
        ("", False),
        ("7. GOVERNING LAW", True),
        ("7.1 This Agreement shall be governed by the laws of India.", False),
        ("7.2 Courts of Bangalore, Karnataka, India shall have exclusive jurisdiction.", False),
        ("", False),
        ("IN WITNESS WHEREOF, the parties have executed this Agreement.", False),
        ("", False),
        ("FOR INNOVATETECH PRIVATE LIMITED:", True),
        ("Ms. Priya Sharma, Head of Human Resources", False),
        ("Date: March 1, 2024", False),
        ("", False),
        ("EMPLOYEE:", True),
        ("Mr. Rajesh Kumar", False),
        ("Date: March 1, 2024", False),
    ]
    
    for text, is_bold in content:
        p = doc.add_paragraph(text)
        if is_bold and text:
            p.runs[0].bold = True
    
    doc.save(filename)
    print(f"[OK] Created: {filename}")

# 3. NDA (TXT)
def create_nda_txt():
    filename = "test_contracts/nda.txt"
    content = """NON-DISCLOSURE AGREEMENT (NDA)

This Non-Disclosure Agreement ("Agreement") is entered into on February 10, 2024, between:

DISCLOSING PARTY: DataSecure Technologies Private Limited
Registered Office: 5th Floor, Cyber Towers, Hitech City, Hyderabad - 500081, Telangana, India
CIN: U72200TG2018PTC125678
("Disclosing Party")

AND

RECEIVING PARTY: CloudVault Systems Private Limited
Registered Office: Unit 12, IT Park, Bandra Kurla Complex, Mumbai - 400051, Maharashtra, India
CIN: U72300MH2019PTC134567
("Receiving Party")

1. DEFINITION OF CONFIDENTIAL INFORMATION

"Confidential Information" means any and all information disclosed by the Disclosing Party to the Receiving Party, including but not limited to:

a) Technical Information: Source code, algorithms, software architecture, technical specifications, research and development projects

b) Business Information: Business plans, strategies, financial information, customer lists, supplier information, marketing plans

c) Product Information: Product roadmaps, unreleased features, specifications, beta versions, prototypes

2. OBLIGATIONS OF THE RECEIVING PARTY

2.1 The Receiving Party agrees to:
- Hold all Confidential Information in strict confidence
- Protect Confidential Information with reasonable care
- Not disclose Confidential Information to any third party without prior written consent
- Use Confidential Information solely for evaluating potential business collaboration

2.2 The Receiving Party may disclose Confidential Information only to employees who have a legitimate need to know and are bound by similar confidentiality obligations.

3. PERMITTED USE

The Receiving Party may use Confidential Information solely for:
- Evaluating potential business collaboration opportunities
- Conducting technical due diligence
- Assessing compatibility of systems and technologies
- Negotiating definitive agreements for business partnership

4. TERM AND SURVIVAL

4.1 Term: This Agreement shall commence on February 10, 2024, and shall continue for 2 years.

4.2 Survival: The confidentiality obligations shall survive for 5 years from the date of disclosure of the Confidential Information.

5. NO LICENSE OR TRANSFER OF RIGHTS

This Agreement does not grant the Receiving Party any license, right, title, or interest in the Confidential Information or any intellectual property rights therein.

6. REMEDIES

The Receiving Party acknowledges that breach of this Agreement may cause irreparable harm and that the Disclosing Party shall be entitled to seek equitable relief, including injunction and specific performance.

7. GOVERNING LAW AND JURISDICTION

7.1 This Agreement shall be governed by the laws of India.

7.2 Courts of Mumbai, Maharashtra, India shall have exclusive jurisdiction.

IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.

FOR DATASECURE TECHNOLOGIES PRIVATE LIMITED:
Mr. Arun Mehta, Chief Executive Officer
Date: February 10, 2024

FOR CLOUDVAULT SYSTEMS PRIVATE LIMITED:
Ms. Kavita Desai, Managing Director
Date: February 10, 2024
"""
    
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"[OK] Created: {filename}")

# 4. LEASE AGREEMENT (PDF)
def create_lease_agreement_pdf():
    filename = "test_contracts/lease_agreement.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, alignment=1)
    story.append(Paragraph("RESIDENTIAL LEASE AGREEMENT", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    content = """
This Lease Agreement ("Agreement") is entered into on April 1, 2024, between:<br/><br/>

<b>LANDLORD:</b> Mr. Suresh Patel<br/>
Address: House No. 15, Sector 21, Dwarka, New Delhi - 110075, India<br/>
("Landlord" or "Lessor")<br/><br/>

<b>AND</b><br/><br/>

<b>TENANT:</b> Ms. Anjali Sharma<br/>
Address: Flat 204, Green Park Extension, New Delhi - 110016, India<br/>
("Tenant" or "Lessee")<br/><br/>

<b>PREMISES:</b> Flat No. 502, Tower A, Skyline Residency, Plot No. 42, Sector 18, Dwarka, New Delhi - 110078<br/><br/>

<b>1. LEASE TERM</b><br/>
1.1 Commencement: The lease shall commence on April 1, 2024.<br/>
1.2 Duration: The initial term shall be for 11 months, ending on February 28, 2025.<br/>
1.3 Renewal: May be renewed upon mutual written agreement 30 days before expiration.<br/><br/>

<b>2. RENT AND PAYMENT TERMS</b><br/>
2.1 Monthly Rent: The Tenant shall pay monthly rent of Thirty Thousand Rupees (₹30,000) per month.<br/>
2.2 Payment Schedule: Rent shall be paid on or before the 5th day of each calendar month.<br/>
2.3 Late Payment: A late fee of ₹500 per day shall be charged for payments after the 5th.<br/><br/>

<b>3. SECURITY DEPOSIT</b><br/>
3.1 Amount: The Tenant shall pay a refundable security deposit of ₹90,000 (three months' rent).<br/>
3.2 Refund: The deposit shall be refunded within 30 days of lease termination, after deducting any amounts due.<br/><br/>

<b>4. USE OF PREMISES</b><br/>
4.1 The Premises shall be used solely for residential purposes.<br/>
4.2 Maximum occupants: 4 persons (Tenant and immediate family).<br/>
4.3 Prohibited: Commercial use, subletting, pets without permission, structural alterations, illegal activities.<br/><br/>

<b>5. MAINTENANCE AND REPAIRS</b><br/>
5.1 Landlord's Responsibilities: Structural integrity, water supply, drainage, common areas.<br/>
5.2 Tenant's Responsibilities: Cleanliness, minor repairs, broken fixtures, reporting major issues.<br/><br/>

<b>6. UTILITIES</b><br/>
Tenant pays: Electricity, water, gas, internet, society maintenance charges.<br/>
Landlord pays: Property taxes, building insurance.<br/><br/>

<b>7. TERMINATION</b><br/>
7.1 Notice Period: Either party may terminate with 2 months written notice.<br/>
7.2 Early Termination: Tenant forfeits one month's rent from security deposit.<br/><br/>

<b>8. GOVERNING LAW</b><br/>
This Agreement shall be governed by the laws of India. Courts of New Delhi shall have exclusive jurisdiction.<br/><br/>

IN WITNESS WHEREOF, the parties have executed this Agreement.<br/><br/>

LANDLORD: Mr. Suresh Patel<br/>
Date: April 1, 2024<br/><br/>

TENANT: Ms. Anjali Sharma<br/>
Date: April 1, 2024
"""
    
    story.append(Paragraph(content, styles['Normal']))
    doc.build(story)
    print(f"[OK] Created: {filename}")

# 5. HINDI SERVICE AGREEMENT (PDF)
def create_hindi_service_agreement_pdf():
    filename = "test_contracts/hindi_service_agreement.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Register Hindi font
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os
    
    # Register Hindi font
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os
    
    hindi_font_name = 'HindiFont'
    
    # Found Nirmala.ttc in C:\Windows\Fonts
    font_path = 'C:\\Windows\\Fonts\\Nirmala.ttc'
    
    try:
        if os.path.exists(font_path):
            # subfontIndex=0 is usually Regular
            pdfmetrics.registerFont(TTFont(hindi_font_name, font_path, subfontIndex=0))
            hindi_style = ParagraphStyle('Hindi', parent=styles['Normal'], fontName=hindi_font_name, fontSize=10)
            title_hindi_style = ParagraphStyle('HindiTitle', parent=styles['Heading1'], fontName=hindi_font_name, fontSize=16, alignment=1)
            print(f"Using Hindi font: {font_path}")
        else:
             # Fallback 
            print("Warning: Nirmala.ttc not found. Trying fallback logic...")
            raise Exception("Nirmala.ttc not found")
            
    except Exception as e:
        print(f"Font registration warning: {e}")
        # Final fallback
        hindi_style = styles['Normal']
        title_hindi_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, alignment=1)
    
    story = []
    
    story.append(Paragraph("सेवा समझौता (SERVICE AGREEMENT)", title_hindi_style))
    story.append(Spacer(1, 0.3*inch))
    
    content = """
यह सेवा समझौता ("समझौता") दिनांक 15 जनवरी 2024 को निम्नलिखित पक्षों के बीच किया गया है:<br/><br/>

<b>सेवा प्रदाता:</b> टेकसॉल्यूशंस प्राइवेट लिमिटेड<br/>
पंजीकृत कार्यालय: प्लॉट नंबर 45, सेक्टर 62, नोएडा - 201301, उत्तर प्रदेश, भारत<br/>
CIN: U72900UP2020PTC123456<br/><br/>

<b>तथा</b><br/><br/>

<b>ग्राहक:</b> श्री विक्रम सिंह<br/>
पता: मकान नंबर 123, सेक्टर 15, गुड़गांव - 122001, हरियाणा, भारत<br/><br/>

<b>1. सेवाएं (SERVICES)</b><br/>
सेवा प्रदाता निम्नलिखित सेवाएं प्रदान करने के लिए सहमत है:<br/>
- ई-कॉमर्स वेबसाइट का विकास<br/>
- Android और iOS के लिए मोबाइल एप्लिकेशन<br/>
- भुगतान गेटवे एकीकरण (Razorpay, Paytm)<br/>
- 3 महीने का निःशुल्क रखरखाव<br/><br/>

<b>2. मुआवजा (COMPENSATION)</b><br/>
2.1 कुल शुल्क: ग्राहक पांच लाख रुपये (₹5,00,000) का भुगतान करेगा।<br/><br/>

2.2 भुगतान अनुसूची:<br/>
- 40% (₹2,00,000) समझौते पर हस्ताक्षर करने पर<br/>
- 30% (₹1,50,000) विकास चरण पूर्ण होने पर<br/>
- 30% (₹1,50,000) अंतिम डिलीवरी पर<br/><br/>

2.3 भुगतान शर्तें: सभी भुगतान चालान से 30 दिनों के भीतर।<br/><br/>

<b>3. अवधि और समाप्ति (TERM)</b><br/>
3.1 अवधि: 15 जनवरी 2024 से 31 दिसंबर 2024 तक।<br/>
3.2 सेवा प्रदाता 15 दिनों की सूचना से समाप्त कर सकता है।<br/>
3.3 ग्राहक 30 दिनों की सूचना से समाप्त कर सकता है।<br/><br/>

<b>4. गोपनीयता (CONFIDENTIALITY)</b><br/>
दोनों पक्ष गोपनीय जानकारी को सुरक्षित रखने के लिए सहमत हैं। गोपनीयता दायित्व समाप्ति के 2 वर्षों तक जीवित रहेंगे।<br/><br/>

<b>5. शासी कानून (GOVERNING LAW)</b><br/>
यह समझौता भारत के कानूनों द्वारा शासित होगा। नोएडा, उत्तर प्रदेश की अदालतों का विशेष अधिकार क्षेत्र होगा।<br/><br/>

साक्षी के रूप में:<br/><br/>

<b>सेवा प्रदाता:</b> टेकसॉल्यूशंस प्राइवेट लिमिटेड<br/>
श्री अमित कुमार, मुख्य कार्यकारी अधिकारी<br/>
दिनांक: 15 जनवरी 2024<br/><br/>

<b>ग्राहक:</b> श्री विक्रम सिंह<br/>
दिनांक: 15 जनवरी 2024
"""
    
    story.append(Paragraph(content, hindi_style))
    doc.build(story)
    print(f"[OK] Created: {filename}")

# Generate all files
if __name__ == "__main__":
    print("Creating test contract files...")
    print()
    
    create_service_agreement_pdf()
    create_employment_contract_docx()
    create_nda_txt()
    create_lease_agreement_pdf()
    create_hindi_service_agreement_pdf()
    
    print()
    print("[SUCCESS] All 5 test files created successfully!")
    print()
    print("Files created:")
    print("1. test_contracts/service_agreement.pdf (Service Agreement - PDF)")
    print("2. test_contracts/employment_contract.docx (Employment Contract - DOCX)")
    print("3. test_contracts/nda.txt (NDA - TXT)")
    print("4. test_contracts/lease_agreement.pdf (Lease Agreement - PDF - Unsupported)")
    print("5. test_contracts/hindi_service_agreement.pdf (Hindi Service Agreement - PDF)")
